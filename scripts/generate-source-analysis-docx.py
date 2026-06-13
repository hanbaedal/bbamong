# -*- coding: utf-8 -*-
"""PPADUN9 소스 코드 분석 보고서 .docx 생성 스크립트"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

OUTPUT = Path(__file__).resolve().parent.parent / "PPADUN9_소스분석.docx"


def set_cell_shading(cell, color_hex: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], "E8E8F0")
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table


def build_document():
    doc = Document()

    # 기본 스타일
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)

    # 표지
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PPADUN9 (빠던9)\n소스 코드 분석 보고서")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1A, 0x6D, 0xFF)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("프로젝트 경로: c:\\PPADUN9\\web")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run("작성일: 2026년 6월 12일").font.size = Pt(11)

    doc.add_page_break()

    # 1. 개요
    add_heading(doc, "1. 프로젝트 개요", 1)
    doc.add_paragraph(
        "이 프로젝트는 야구 경기 실시간 예측·포인트·커뮤니티 플랫폼인 PPADUN9(빠던9)입니다. "
        "웹과 모바일(Android/iOS)을 함께 지원하며, 일반 사용자 / 관리자 / 매니저 세 가지 앱이 "
        "하나의 코드베이스에서 동작하는 풀스택 모노레포 구조입니다."
    )

    add_heading(doc, "1.1 핵심 가치", 2)
    for item in [
        "실시간 라운드 예측 — WebSocket 기반 동기화, 매니저가 현장에서 결과 입력",
        "포인트·기부·랭킹 — 예측 성공 시 풀 분배, 기부, 승리/포인트 랭킹",
        "3역할 분리 — 사용자 / 어드민 / 매니저가 각각 독립 UI·인증·세션 정책",
        "모바일 네이티브 — Capacitor로 iOS/Android 앱, OAuth 딥링크, AdMob",
        "운영 자동화 — 비활성 로그아웃, 경기 자동 종료, 정지 회원 정리 배치",
    ]:
        add_bullet(doc, item)

    # 2. 디렉터리 구조
    add_heading(doc, "2. 전체 디렉터리 구조", 1)
    structure = """web/
├── client/              # React 프론트엔드 (User / Admin / Manager)
├── server/              # Express 백엔드 API + WebSocket
├── shared/              # Drizzle ORM 스키마 (프론트·백 공유)
├── assets/              # 이미지·SVG 등 정적 리소스
├── android-manager/     # 매니저 앱 Android 네이티브 셸
├── ios-manager-standalone/  # 매니저 iOS
├── capacitor.config.ts      # 사용자 앱 Capacitor 설정
└── capacitor.config.manager.ts  # 매니저 앱 설정"""
    p = doc.add_paragraph()
    run = p.add_run(structure)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    add_heading(doc, "2.1 빌드·실행 명령", 2)
    add_table(
        doc,
        ["명령", "역할"],
        [
            ["npm run dev", "Vite(프론트) + Express(백) 개발 서버"],
            ["npm run build", "Vite 빌드 + esbuild로 서버 번들"],
            ["npm run start", "프로덕션 서버 실행"],
            ["npm run db:push", "Drizzle로 DB 스키마 반영"],
        ],
    )

    # 3. 3개 앱
    add_heading(doc, "3. 3개 앱 구조 (단일 진입점 분기)", 1)
    doc.add_paragraph(
        "client/src/main.tsx에서 URL 경로에 따라 앱을 분기합니다. "
        "/manager로 시작하면 ManagerApp, / 또는 /admin이면 AdminApp, "
        "그 외는 UserApp이 렌더링됩니다."
    )
    add_table(
        doc,
        ["앱", "경로", "대상", "주요 기능"],
        [
            ["UserApp", "/login, /home, /prediction 등", "일반 회원", "예측, 출석, 게시판, 포인트, 설정"],
            ["AdminApp", "/, /admin/*", "슈퍼/일반 어드민", "회원·경기·광고·모니터링 관리"],
            ["ManagerApp", "/manager/*", "현장 매니저", "경기 결과 입력, 실시간 운영"],
        ],
    )
    doc.add_paragraph(
        "각 앱은 별도 QueryClient, 토큰 관리, 에셋 프리로더를 사용하여 역할별로 격리되어 있습니다."
    )

    # 4. 사용자 앱
    add_heading(doc, "4. 사용자 앱 (UserApp)", 1)
    add_heading(doc, "4.1 인증", 2)
    for item in [
        "로그인/회원가입, 비밀번호 찾기",
        "소셜 로그인: Kakao, Google, Apple",
        "게스트 로그인 (프로필, 고객센터, 게시글 작성 등 제한)",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4.2 주요 화면", 2)
    add_table(
        doc,
        ["경로", "기능"],
        [
            ["/home", "홈 (기본 로그인 후 랜딩)"],
            ["/prediction", "실시간 경기 예측 (WebSocket)"],
            ["/attendance", "출석 체크"],
            ["/board, /create-post, /post-detail", "커뮤니티 게시판"],
            ["/point, /point-history", "포인트 관리"],
            ["/setting/*", "프로필, 초대, 기부 내역, FAQ, 공지, 전자책 등"],
        ],
    )

    add_heading(doc, "4.3 모바일 특화", 2)
    for item in [
        "Capacitor 기반 네이티브 앱 (com.bbanden.nine, 앱명 '빠던9')",
        "Android 뒤로가기: /home, /login에서만 앱 최소화, 나머지는 history.back()",
        "AdMob 광고 (useAdMob.ts)",
        "Safe area, 375px 모바일 퍼스트 UI",
    ]:
        add_bullet(doc, item)

    # 5. 관리자 앱
    add_heading(doc, "5. 관리자 앱 (AdminApp)", 1)
    doc.add_paragraph("루트(/)가 어드민 로그인으로 연결됩니다.")
    add_heading(doc, "5.1 메뉴 구성 (AdminSidebar)", 2)
    for item in [
        "회원 관리 — 회원 리스트, 사회공헌참여기록 관리",
        "운영자 관리 — 직원 리스트(슈퍼어드민), 운영자 리스트, 경기 할당, 상태 모니터링",
        "수익 관리 — 동영상 광고 수익 현황",
        "경기 관리 — 경기 생성/관리",
        "공지사항, 고객 지원 센터, 약관 관리",
    ]:
        add_bullet(doc, item)
    doc.add_paragraph(
        "어드민은 승인 대기 → 승인 흐름이 있으며, 로그인 시 선점 방식(이미 로그인 중이면 409)으로 "
        "중복 로그인을 방지합니다."
    )

    # 6. 매니저 앱
    add_heading(doc, "6. 매니저 앱 (ManagerApp)", 1)
    for item in [
        "로그인/회원가입/승인 대기",
        "홈: 배정된 경기 목록",
        "matchDetail: 라운드별 결과 입력·예측 시작/종료 제어",
        "Android/iOS 별도 빌드 (android-manager/, ios-manager-standalone/)",
        "Error Boundary로 크래시 시 로그인 화면 복귀",
    ]:
        add_bullet(doc, item)

    # 7. 백엔드
    add_heading(doc, "7. 백엔드 아키텍처", 1)
    add_heading(doc, "7.1 API 라우트 구조", 2)
    for item in [
        "/api/users, /api/points, /api/posts ... → 일반 사용자 API",
        "/api/admin/* → 관리자 API",
        "/api/manager/* → 매니저 API",
        "/api/live-match/* → 실시간 예측 API",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7.2 인증 미들웨어", 2)
    add_table(
        doc,
        ["파일", "역할"],
        [
            ["userAuth.ts", "일반 사용자 JWT"],
            ["adminAuth.ts", "어드민 (역할·승인 상태 검증)"],
            ["managerAuth.ts", "매니저"],
            ["multiRoleAuth.ts", "복수 역할"],
        ],
    )

    add_heading(doc, "7.3 세션·인증 정책", 2)
    for item in [
        "JWT + Refresh Token Rotation",
        "토큰: httpOnly, secure, sameSite=strict 쿠키",
        "세션 저장소: Redis (sessionManager.ts)",
        "일반 사용자: 새 기기 로그인 시 기존 세션 강제 종료 (force-replace)",
        "어드민/매니저: 첫 로그인 우선 (409 차단)",
        "WS 종료 코드 4005(세션 종료) 시 자동 재연결",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7.4 실시간 경기 시스템 (server/liveMatch/)", 2)
    add_table(
        doc,
        ["모듈", "역할"],
        [
            ["wsManager.ts", "WebSocket 연결·인증·브로드캐스트"],
            ["broadcastManager.ts", "이벤트 전파"],
            ["predictionRoutes.ts", "예측 제출 API"],
            ["matchControlRoutes.ts", "라운드 시작/종료/결과 (매니저·어드민)"],
            ["predictionStorage.ts", "DB 연동"],
        ],
    )

    doc.add_paragraph("WebSocket 이벤트 예시:")
    for evt in [
        "prediction_started, prediction_ended",
        "round_result, round_next",
        "stats_update, ad_started, ad_stopped",
        "waiting_screen_update, match_end",
    ]:
        add_bullet(doc, evt)

    doc.add_paragraph(
        "예측은 라운드 단위(1루/2루/3루/홈런/아웃)이며, 풀 기반 보상·기부 시스템이 적용됩니다."
    )

    add_heading(doc, "7.5 배치 작업", 2)
    add_table(
        doc,
        ["파일", "역할"],
        [
            ["inactiveLogoutBatch.ts", "비활성 사용자 자동 로그아웃 (WS 연결·5분 유예 제외)"],
            ["matchAutoCloseBatch.ts", "KST 기준 지난 경기 자동 완료 (시작 시 + 매시간)"],
            ["suspendedUserCleanupBatch.ts", "정지 7일 경과 회원 영구 삭제"],
        ],
    )

    # 8. DB
    add_heading(doc, "8. 데이터베이스 스키마 (shared/schema.ts)", 1)
    doc.add_paragraph("Drizzle ORM + Zod 검증 스키마를 프론트엔드와 백엔드가 공유합니다.")
    add_table(
        doc,
        ["테이블", "설명"],
        [
            ["users", "회원 (소셜/게스트, 초대코드, 포인트, 정지 상태)"],
            ["matches", "경기 (구장, 일시, 상태, 현재 라운드)"],
            ["predictions", "사용자별 라운드 예측"],
            ["round_statistics", "라운드별 참가·포인트·승자 통계"],
            ["point_transactions", "포인트 적립/사용 내역"],
            ["attendance_records", "출석 기록"],
            ["posts, comments", "게시판"],
            ["admin_users", "어드민/매니저 계정"],
            ["advertisements, waiting_screens", "광고·대기 화면"],
            ["inquiries, notices, terms, faqs", "고객지원·공지"],
            ["ebooks, ebook_purchases", "전자책"],
        ],
    )
    doc.add_paragraph("경기 상태: scheduled → ongoing → completed / cancelled")

    # 9. 프론트엔드 기술 스택
    add_heading(doc, "9. 프론트엔드 기술 스택", 1)
    add_table(
        doc,
        ["영역", "기술"],
        [
            ["UI", "shadcn/ui (Radix UI), Tailwind CSS"],
            ["라우팅", "Wouter"],
            ["서버 상태", "TanStack Query"],
            ["폼", "React Hook Form + Zod"],
            ["애니메이션", "Framer Motion, Lottie"],
            ["차트", "Recharts (어드민)"],
            ["파일 업로드", "Uppy + Replit Object Storage"],
        ],
    )

    # 10. 인프라
    add_heading(doc, "10. 인프라·배포", 1)
    for item in [
        "Replit 환경 기준 (replit.md, replit.nix)",
        "DB: Neon (서버리스 PostgreSQL)",
        "세션: Redis (서버 시작 시 자동 기동 시도)",
        "파일: Replit Object Storage / Google Cloud Storage",
        "프로덕션 도메인: https://ppamong.com",
        "Capacitor 앱은 원격 URL 로드 (capacitor.config.ts의 server.url)",
    ]:
        add_bullet(doc, item)

    # 11. 아키텍처 요약
    add_heading(doc, "11. 시스템 아키텍처 요약", 1)
    arch = """[클라이언트]
  User App (웹/모바일) ──┐
  Admin App (웹)        ──┼──→ Express REST API ──→ PostgreSQL (Neon)
  Manager App (Android/iOS)─┘         │
                                      ├──→ Redis (세션)
  모든 클라이언트 ──→ WebSocket Server ──┘
  배치 작업 ──→ PostgreSQL
  Express ──→ Object Storage (파일/영상)"""
    p = doc.add_paragraph()
    run = p.add_run(arch)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    # 12. 백엔드 기술 스택
    add_heading(doc, "12. 백엔드 기술 스택", 1)
    add_table(
        doc,
        ["영역", "기술"],
        [
            ["런타임", "Node.js + TypeScript + Express"],
            ["ORM", "Drizzle ORM"],
            ["DB", "PostgreSQL (Neon)"],
            ["세션", "Redis (ioredis)"],
            ["인증", "JWT, bcrypt, Passport"],
            ["실시간", "WebSocket (ws)"],
            ["검증", "Zod"],
            ["빌드", "esbuild (서버), Vite (클라이언트)"],
        ],
    )

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("— 보고서 끝 —")
    fr.font.size = Pt(10)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    return doc


def main():
    doc = build_document()
    doc.save(str(OUTPUT))
    print(f"생성 완료: {OUTPUT}")


if __name__ == "__main__":
    main()
