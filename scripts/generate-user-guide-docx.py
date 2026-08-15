# -*- coding: utf-8 -*-
"""빠몽이 사용 설명서 .docx 생성 — 앱 UserGuideContent 기준"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent.parent / "docs" / "빠몽이_사용설명서.docx"

# shared/predictionOdds.ts 와 동기화
PREDICTION_ODDS = {"아웃": 1.2, "1루": 1.5, "2루": 3, "3루": 10, "홈런": 5}
BET_AMOUNT_OPTIONS = [50, 100, 200, 500, 1000]
SIDE_BET_AMOUNT_OPTIONS = [100, 200, 300, 500, 700, 1000]
WINNER_ODDS = 2
EXACT_SCORE_ODDS = 20
AD_REWARD_POINTS = 500
AD_EARLY_DISMISS_SECONDS = 5

SECTIONS: list[tuple[str, list[str]]] = [
    (
        "앱 시작·로그인",
        [
            "앱을 완전히 종료했다가 다시 열면, 저장된 로그인(세션)을 확인합니다.",
            "세션이 있으면 인트로 없이 바로 홈으로 이동합니다.",
            "세션이 없으면 환영 인트로(약 3.5초) 후 로그인 화면이 나옵니다. 로그인 화면 왼쪽에 이용 안내(15세·재화)가 표시됩니다.",
            "로그인: 회원 아이디·비밀번호, 소셜(카카오·구글·애플), 또는 게스트로 시작할 수 있습니다.",
            "한 계정은 한 기기에만 로그인됩니다. 다른 기기에서 이미 로그인 중이면 안내 메시지가 표시됩니다.",
            "다른 앱을 보다가 돌아올 때(백그라운드 복귀)는 로그인 화면 없이 마지막 화면을 이어갑니다.",
        ],
    ),
    (
        "홈 화면",
        [
            "가로 화면 기준으로 왼쪽·오른쪽 두 영역으로 나뉩니다.",
            "왼쪽: 빠몽이 캐릭터(탭 가능)와 「예측게임 하러가기」 버튼 — 둘 다 예측 게임으로 이동합니다.",
            "오른쪽: 인사말과 메뉴 — 「야구 예측 게임이란?」, 「사용설명서」, 「공지사항」, 「문의하기」, 「게시판」, 「빠몽이 쇼핑센터」 등",
            "「야구 예측 게임이란?」은 왼쪽 패널에서, 「사용설명서」는 안내 창에서 열립니다.",
            "공지·문의·게시판은 각각 해당 화면으로 이동합니다.",
            "홈 화면에는 하단 메뉴가 없습니다.",
        ],
    ),
    (
        "예측 화면 진입",
        [
            "홈 왼쪽의 빠몽이 또는 「예측게임 하러가기」로 예측 화면(/prediction)에 들어갑니다.",
            "경기 시작 1분 전 이전(또는 타석 참여 가능 시간이 아닐 때)에는 「오늘의 경기」 모달이 자동으로 열립니다.",
            "모달에는 DB에 등록된 오늘 경기(최대 5경기)가 표시됩니다.",
            "경기가 시작되어 타석 예측이 가능한 시간대에는 모달이 자동으로 뜨지 않습니다.",
            "상단 경기명·경기장을 눌러 참여 가능한 경기·경기장을 바꿀 수 있습니다.",
        ],
    ),
    (
        "예측 게임 (타석)",
        [
            "타석 예측은 경기 시작 1분 전부터 종료 전까지 가능합니다.",
            "예측이 열리면 아웃·1루·2루·3루·홈런 중 하나를 고릅니다. (실황 자동 또는 운영자 「예측 시작」)",
            "「1루」는 1루타·포볼·데드볼 등 1루 진루 결과를 포함합니다.",
            f"배팅 포인트({', '.join(str(x) for x in BET_AMOUNT_OPTIONS)}P)를 선택한 뒤 확인하면 즉시 차감됩니다.",
            "적중 시 선택금액 × 고정배당이 지급되고, 미적중 시 배팅분은 소멸합니다.",
            "타석 결과가 확정되면 적중/미적중 연출 후 다음 타석을 기다립니다. (실황 자동 확정 또는 운영자 입력)",
            "투수 교체 등으로 진행 중이던 예측이 취소되면, 해당 배팅은 환불될 수 있습니다.",
            "경기 종료 시 약 10초 「경기종료」 안내가 표시된 뒤 홈으로 이동합니다.",
        ],
    ),
    (
        "승리팀 · 최종 스코어",
        [
            "「오늘의 경기」 모달에서 경기마다 「우승팀 맞추기」「점수 맞추기」로 배팅합니다.",
            "버튼이 활성(클릭 가능)인 경기만 배팅할 수 있습니다.",
            "팀 표시는 「홈팀」「원정팀」만 사용합니다 (구단명은 표시하지 않음).",
            f"배팅 금액: {', '.join(str(x) for x in SIDE_BET_AMOUNT_OPTIONS)}P",
            f"승리팀 맞추기: {WINNER_ODDS}배 (100P → {100 * WINNER_ODDS}P)",
            f"최종 스코어 정확히: {EXACT_SCORE_ODDS}배 (100P → {100 * EXACT_SCORE_ODDS}P)",
            "1회 시작 시 자동 마감 — 이후 신규·변경 불가 (마감된 경기는 버튼 비활성)",
            "경기 종료 후 실황 최종 스코어로 자동 정산 · 모달에 적중/미적중/환불 표시",
            "경기 취소·무승부 시 해당 배팅은 환불",
        ],
    ),
    (
        "배당표 (타석)",
        [
            f"{k}: {v}배 (예: 100P 적중 → {int(100 * v)}P)"
            for k, v in PREDICTION_ODDS.items()
        ],
    ),
    (
        "광고·보상",
        [
            "공수교대·투수교체 때 전면(및 보상형) 광고가 나올 수 있습니다. 예측 게임 중 하단 배너 광고는 없습니다.",
            f"광고가 시작된 뒤 약 {AD_EARLY_DISMISS_SECONDS}초가 지나면 「×」로 끌 수 있습니다. 끄면 보상은 없습니다.",
            f"운영자가 광고를 중지할 때까지 시청하면 {AD_REWARD_POINTS}P가 지급됩니다. (짧은 자동 종료만으로는 보상 없음)",
            "「예측 시작」으로 광고가 중지되거나, 5초 만에 「×」로 끄면 보상은 없습니다.",
        ],
    ),
    (
        "예측 화면 메뉴",
        [
            "왼쪽 세로 메뉴: 「홈」, 「내이야기」, 「쇼핑센터」, 「내정보」",
            "내이야기: 승리현황, 친구 초대, 출석 체크, 나의 콘텐츠, 사회공헌 참여현황",
            "내정보: 회원정보, 추가 참여, Q&A, 서비스 이용약관, 탈퇴하기",
            "「쇼핑센터」: 빠몽이 쇼핑센터로 이동합니다.",
        ],
    ),
    (
        "헤더·기타",
        [
            "가운데 로고: 홈으로 이동",
            "홈 우측 상단: 로그아웃",
            "화면 하단에는 사이드 배팅(우승팀·점수) 요약이 표시될 수 있습니다.",
            "예측 화면 오른쪽 상단 공지 배지를 한 번 닫으면, 같은 계정에서는 해당 공지가 다시 표시되지 않습니다.",
        ],
    ),
    (
        "연습 팁",
        [
            "게임 소개는 홈의 「야구 예측 게임이란?」에서 확인하세요.",
            "「게임 시뮬레이션」에서 예측 화면·내이야기·내정보 안내와 사이드·타석·정산 흐름을 연습하세요. 왼쪽 단계 탭으로 건너뛸 수 있습니다.",
            "시뮬레이션은 연습용이며 보유 포인트에 영향이 없습니다.",
            f"타석 선택 금액: {', '.join(str(x) for x in BET_AMOUNT_OPTIONS)}P · 사이드: {', '.join(str(x) for x in SIDE_BET_AMOUNT_OPTIONS)}P",
        ],
    ),
]


def set_run_font(run, *, name: str = "맑은 고딕", size: int | None = None, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11)


def main() -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    title = doc.add_heading("빠몽이 사용 설명서", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_run_font(run, size=22, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("PPAMONG (ppamong.com) · 사용자 앱 안내")
    set_run_font(r, size=12)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    intro = doc.add_paragraph()
    r = intro.add_run(
        "빠몽이 앱 사용법·메뉴·게임 흐름을 안내합니다. "
        "게임 소개는 홈의 「야구 예측 게임이란?」을 참고하세요. "
        "본 문서는 앱 내 「사용설명서」와 동일한 구성을 따릅니다."
    )
    set_run_font(r, size=11)

    doc.add_paragraph()

    for idx, (section_title, items) in enumerate(SECTIONS, start=1):
        h = doc.add_heading(f"{idx}. {section_title}", level=1)
        for run in h.runs:
            set_run_font(run, size=14, bold=True)
        add_bullets(doc, items)

    doc.add_heading("부록. 타석 배당 한눈에 보기", level=1)
    table = doc.add_table(rows=1 + len(PREDICTION_ODDS), cols=3)
    table.style = "Table Grid"
    headers = ["예측", "배당", "100P 적중 시"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, size=11, bold=True)
    for ri, (k, v) in enumerate(PREDICTION_ODDS.items(), start=1):
        vals = [k, f"{v}배", f"{int(100 * v)}P"]
        for ci, val in enumerate(vals):
            cell = table.rows[ri].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=11)

    doc.add_paragraph()
    note = doc.add_paragraph()
    r = note.add_run(
        "※ 배당·배팅 금액·광고 보상은 서비스 정책에 따라 변경될 수 있습니다. "
        "최신 내용은 앱 내 사용설명서를 확인해 주세요."
    )
    set_run_font(r, size=9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
