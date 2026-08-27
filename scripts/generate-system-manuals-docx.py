# -*- coding: utf-8 -*-
"""시스템 매뉴얼 DOCX 일괄 생성 (관리자·쇼핑몰·운영자·DB). 사용자용은 generate-user-guide-docx.py"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"


def set_run_font(run, *, size: int | None = None, bold: bool = False) -> None:
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold


def new_doc(title: str, subtitle: str) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        set_run_font(run, size=20, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(subtitle)
    set_run_font(r, size=11)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()
    return doc


def add_section(doc: Document, title: str, bullets: list[str]) -> None:
    h = doc.add_heading(title, level=1)
    for run in h.runs:
        set_run_font(run, size=14, bold=True)
    for item in bullets:
        p = doc.add_paragraph(style="List Bullet")
        set_run_font(p.add_run(item), size=11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, size=11, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=11)
    doc.add_paragraph()


def footer_note(doc: Document) -> None:
    p = doc.add_paragraph()
    r = p.add_run(
        "※ 본 문서는 GitHub docs/ 원본을 기준으로 생성됩니다. "
        "최신본은 관리자 「시스템 매뉴얼」에서 다운로드하세요."
    )
    set_run_font(r, size=9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def write_admin() -> Path:
    doc = new_doc(
        "빠몽이 사용 설명서 — 관리자",
        "PPAMONG Admin (/admin) · 2026-08-27 운영 기준",
    )
    add_section(
        doc,
        "1. 개요",
        [
            "관리자 웹은 회원·경기·운영자·수익·고객지원·몰을 총괄합니다.",
            "슈퍼바이저만 관리자 등록·시스템 운영·시스템 매뉴얼에 접근합니다.",
            "로그인: /admin/login — 승인된 관리자만 사용합니다.",
            "최신 운영 요약은 관리자 「시스템 매뉴얼」 페이지(/admin/ops/system-manuals)와 동일합니다.",
        ],
    )
    add_section(
        doc,
        "2. 주요 메뉴 (현재 사이드바)",
        [
            "기본: 앱 홈 설정 · KBO 선수단 · 오늘의 선발명단 · 앱 파일 등록/다운로드",
            "쇼핑몰·판매: 쇼핑몰 확인·관리 · 주문·판매·재고·구매",
            "슈퍼바이저: 관리자 등록/리스트 · 시스템 매뉴얼 · 디비 백업 · 관리자·운영자 로그인 현황",
            "수익: 동영상 광고 수익·관리 · 배너 수익 · 대기 화면",
            "경기·회원: 경기 관리(달력) · 실시간 게임 모니터링 · 운영자 등록·리스트 · 회원·랭킹·초대",
            "고객 지원: 공지 · 회원 문의 · 게시판 · 약관",
        ],
    )
    add_section(
        doc,
        "3. 일일 체크리스트",
        [
            "경기 관리에서 오늘 KBO 일정을 등록·저장합니다. 팀 로고는 다음 스포츠입니다.",
            "「오늘의 선발명단」으로 타순을 맞춥니다(다음으로 경기 찾기, 네이버로 타순).",
            "운영자 리스트에서 해당 경기 「실황 ON」(다음+네이버+회원 게임). 기본 1경기만 ON.",
            "실시간 게임 모니터링에서 배팅 분포·사이드벳·스코어를 확인합니다.",
            "점수 이상이 있으면 수동 보정 후, 끝나면 수동을 끄고 실황에 맡깁니다.",
            "경기 종료 후 사이드벳 정산·문의·공지를 확인합니다.",
        ],
    )
    add_section(
        doc,
        "4. 실황·예측 운영 포인트",
        [
            "KBO 일정·점수·이닝·로고는 다음 스포츠, 주자·B-S·OUT·타자·구종·상대전적은 네이버입니다. API-SPORTS 키는 필요 없습니다. 같은 필드를 두 소스에서 섞지 않습니다.",
            "「예측 시작」은 predictionEnabled만 켭니다. matchStatus=ongoing은 다음 실황 근거로만 올립니다. 시작 전(BEFORE/READY)이면 scheduled로 되돌립니다.",
            "스코어 PATCH는 controlMode=manual. 「수동」을 끄면 auto로 돌아가 다음 점수를 다시 받습니다.",
            "운영은 하이브리드만입니다. 운영자 UI에 실황 자동 ON/OFF 토글은 없습니다. 리스트의 「실황 ON」은 다음·네이버 폴링 + 회원 게임 연동입니다.",
            "타석 참여는 경기 시작 5분 전부터입니다. 예측 창은 열린 뒤 약 8초 후 자동 중지될 수 있습니다.",
            "타자명 안정화 약 2초, 투수명 기본 6초(최소 3초), 실황 폴링 기본 2초.",
            "공수교대·투수교체 = 리워드 광고 50초. 예측 시작으로 끄면 보상 없음. 게임 중 하단 배너 광고는 없습니다.",
        ],
    )
    add_table(
        doc,
        ["항목", "소스", "비고"],
        [
            ["득점·이닝표·로고", "다음 스포츠", "controlMode=auto일 때 점수 덮어씀"],
            ["주자·B-S·OUT·타자", "네이버 문자중계", "수동 점수여도 주자는 네이버"],
            ["상대전적", "네이버 preview", "스코어보드 아래"],
            ["선발명단 타순", "네이버", "경기는 다음으로 찾음"],
        ],
    )
    add_section(
        doc,
        "5. 회원 예측 화면 (한 타석)",
        [
            "선택 화면(3D 구장)과 주루(필리스 실사)는 베이스 좌표가 다릅니다.",
            "실패·투수교체·공수교대는 3D 구장을 유지합니다. 홈런 주루는 1·2·3루를 돌아 홈(중견으로 가지 않음).",
        ],
    )
    add_table(
        doc,
        ["순서", "단계", "배경", "화면에서 하는 일"],
        [
            ["1", "경기전", "쿠어스 전경", "시작 카운트다운"],
            ["2", "대기", "시네마틱 빠몽이 (초/말)", "다음 타자 대기"],
            ["3", "예측 선택", "3D 빈 구장", "베이스 버튼·포인트"],
            ["4", "결과 대기", "시네마틱 투수 (초/말)", "내 예측 배지"],
            ["5", "결과 글씨", "4와 같음", "약 2.2초 큰 글씨"],
            ["6", "주루(적중)", "필리스 실사", "홈런은 1·2·3루를 돌아 홈"],
            ["7", "다음 타석", "대기 또는 선택", "축하 점프 생략"],
        ],
    )
    add_table(
        doc,
        ["예측", "배당", "100P 적중"],
        [
            ["아웃", "1.2배", "120P"],
            ["1루", "1.5배", "150P"],
            ["2루", "3배", "300P"],
            ["3루", "10배", "1000P"],
            ["홈런", "5배", "500P"],
        ],
    )
    footer_note(doc)
    out = DOCS / "빠몽이_사용설명서_관리자.docx"
    doc.save(out)
    return out


def write_operator() -> Path:
    doc = new_doc("빠몽이 사용 설명서 — 운영자", "PPAMONG Manager · 하이브리드 실황 · 2026-08-27")
    add_section(
        doc,
        "1. 로그인·세션",
        [
            "카카오톡 로그인 링크로만 접속합니다 (앱/웹).",
            "같은 운영자 계정은 한 기기만 사용 가능합니다.",
            "경기 종료 시 약 10초 「경기종료」 안내 후 로그아웃됩니다(「세션 만료」가 아님).",
        ],
    )
    add_section(
        doc,
        "2. 하이브리드 실황 (토글 없음)",
        [
            "실황이 타석을 진행하고, 운영자가 버튼을 먼저 누르면 그게 우선입니다. 완전자동/반자동 토글은 없습니다.",
            "타석 상태머신: 대기 → 예측열림 → 예측닫힘 → 결과확정 → 다음타자/공수교대.",
            "가드: 타자명 약 2초·투수명 기본 6초(최소 3초) 안정화. 예측 열림 기본 8초 후 자동 중지.",
            "결과 확정 전에는 다음타자·공수교대·광고·투수교체를 막습니다.",
            "자동 결과: 아웃(아웃수↑)·1~3루·홈런(아웃 유지+타자 교체). 희생/병살→아웃, 야수선택→1루. 애매하면 「실황 추정」만 하고 1탭 확정.",
            "실황 타자 ≠ 선발이면 대타 표시·자동 pinch. 투수교체는 대타를 유지합니다.",
            "예전 liveAutoEnabled=false 잔여값은 폴링 시 true로 복구됩니다.",
        ],
    )
    add_section(
        doc,
        "3. 수동 컨트롤 (예외·보정)",
        [
            "경기전(scheduled)에는 예측시작·중지·결과·다음타자·투수교체·공수교대·대타가 비활성입니다.",
            "경기중(ongoing)에만 위 버튼을 사용합니다.",
            "수동 정상 흐름: 예측시작 → 예측중지 → 결과 전송 → 「다음 타자」(3아웃이면 「공수교대」).",
            "결과 전송 후 다음 타자/공수교대를 눌러야 다음 타석으로 넘어갑니다(자동 OFF이거나 가드에 걸린 경우).",
            "투수교체·공수교대 = 광고 시작 / 「예측 시작」 = 광고 중지(보상 없음). 운영자가 광고를 중지하면 500P. 별도 「광고 시작」 버튼 없음. 광고 재생은 50초.",
            "투수교체는 같은 타석이므로 대타(pinch)를 지우지 않습니다. 예측 중 투수교체 시 배팅은 환불·결과 생략될 수 있습니다.",
            "대타: 실황 타자명이 선발과 다르면 「대타가 나옵니다」 표시·자동 pinch. 수동 입력도 가능.",
            "팀명 클릭: 주전 1~9 타순·시즌 전적 입력(라인업 피커).",
            "점수 보정은 수동 모드로 잠급니다. 이닝 표시·팀 점수는 실황 스코어보드를 우선합니다.",
        ],
    )
    add_section(
        doc,
        "4. 상태 표시",
        [
            "구장명 · 경기전 / N회 초·말 / 경기종료 / 연기됨 · 타석 단계 배지",
            "네트워크 끊김 시 「실시간 연결 재시도 중」 — 예측 시작·중지는 HTTP로 계속 사용 가능합니다.",
            "회원이 보는 화면: 경기전(쿠어스) → 대기(시네마틱) → 선택(3D 구장) → 결과대기(시네마틱) → 큰 글씨 → 적중 주루(실사, 홈런은 1·2·3루 후 홈).",
        ],
    )
    footer_note(doc)
    out = DOCS / "빠몽이_사용설명서_운영자.docx"
    doc.save(out)
    return out


def write_mall() -> Path:
    doc = new_doc("빠몽이 사용 설명서 — 쇼핑몰", "빠몽이 쇼핑센터 · 몰 관리")
    add_section(
        doc,
        "1. 사용자 화면",
        [
            "운영 URL: https://ppamong.com/shop (회원가입은 사용자 앱만, 몰에 가입 폼 없음).",
            "홈 메뉴 「빠몽이 쇼핑센터」 또는 예측 화면 「쇼핑센터」로 이동합니다.",
            "정회원(게스트 아님)만 주문. 게스트·비로그인은 둘러보기·장바구니만.",
            "1차 결제는 현금 주문 접수. 게임 포인트로 직접 결제하지 않습니다.",
        ],
    )
    add_section(
        doc,
        "2. 관리자 몰 메뉴",
        [
            "몰 미리보기·홈/상품 관리",
            "주문 관리·매출·재고·매입",
            "앱 홈 설정(쇼핑 노출 문구 등)",
        ],
    )
    add_section(
        doc,
        "3. 정책 요약",
        [
            "판매 유형·포인트/결제 정책은 docs/PPAMONG_몰_정책.md, PPAMONG_몰_판매유형.md를 참고합니다.",
            "상품 이미지는 R2 등 오브젝트 스토리지를 사용할 수 있습니다.",
        ],
    )
    footer_note(doc)
    out = DOCS / "빠몽이_사용설명서_쇼핑몰.docx"
    doc.save(out)
    return out


def write_db() -> Path:
    doc = new_doc("빠몽이 DB 구조 설명서", "운영 DB = MongoDB 전용")
    add_section(
        doc,
        "1. 개요",
        [
            "빠몽이(PPAMONG)의 운영·런타임 데이터베이스는 MongoDB만 사용합니다.",
            "서버는 MONGODB_URI 없이 기동하지 않습니다. 회원·경기·예측·몰·관리자 데이터는 모두 MongoDB 컬렉션에 저장됩니다.",
            "구제품(빠던9 등)이 쓰던 PostgreSQL은 빠몽이 운영 DB가 아닙니다. 일상 운영·백업·복구는 MongoDB 기준입니다.",
            "컬렉션 관계는 docs/db-erd.md(구 ERD는 참고용)와 본 문서의 표를 함께 보세요.",
        ],
    )
    add_table(
        doc,
        ["영역", "주요 컬렉션/모델", "설명"],
        [
            ["회원", "User, AttendanceRecord, PointTransaction", "계정·출석·포인트"],
            ["경기", "Match, Stadium, KboPlayer", "일정·스코어보드·타순·대타"],
            ["예측", "Prediction, RoundStatistics, MatchSideBet", "타석·라운드·사이드벳"],
            ["운영", "AdminUser", "스태프·슈퍼바이저·운영자"],
            ["콘텐츠", "Notice, Inquiry, Post, Term", "공지·문의·게시판·약관"],
            ["광고", "Advertisement, WaitingScreen, AppAdmobConfig", "배너·영상·대기·AdMob"],
            ["몰", "GoodsProduct, MallOrder, MallStock…", "상품·주문·재고"],
            ["소셜", "FriendRoom", "친구·동호회 방(공개 예측 함께 참여)"],
        ],
    )
    add_section(
        doc,
        "2. Match 운영 필드 (요약)",
        [
            "matchStatus: scheduled | ongoing | completed | cancelled",
            "gameInning / inningHalf / batterIndexInHalf / outsInHalf — 운영자 진행 기준",
            "matchLineup / matchPlayerStats — 주전 타순·시즌 스탯",
            "pinchHitter — 현재 타석 대타 (다음 타자·공수교대 시 해제, 투수교체 시 유지)",
            "sideBetsLocked / predictionEnabled — 사이드벳 마감·타석 예측 오픈",
            "liveScoreboard / controlMode — 다음 스포츠 실황·수동 스코어 (auto|manual)",
            "liveAutoEnabled — 기본 true. 운영자 UI 토글 없음. false 잔여는 폴링 시 복구",
            "atBatPhase — 타석 상태머신(WS at_bat_phase). 회원 화면 uiStage의 권위",
            "daumGameId — 다음 스포츠 경기 ID. apiSportsGameId는 레거시",
        ],
    )
    add_section(
        doc,
        "3. 백업",
        [
            "슈퍼바이저 → 시스템 운영 → 디비 백업하기",
            "MongoDB 컬렉션/테이블별 내보내기(백업)를 사용합니다.",
            "관리자 화면에 남아 있는 PostgreSQL 가져오기 메뉴는 구제품 일회성 이전용이며, 빠몽이 정상 운영에 필요하지 않습니다.",
        ],
    )
    footer_note(doc)
    out = DOCS / "빠몽이_DB구조_설명서.docx"
    doc.save(out)
    return out


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    paths = [write_admin(), write_operator(), write_mall(), write_db()]
    # 사용자용은 별도 스크립트 유지
    user_script = ROOT / "scripts" / "generate-user-guide-docx.py"
    if user_script.exists():
        import runpy

        runpy.run_path(str(user_script), run_name="__main__")
        paths.append(DOCS / "빠몽이_사용설명서.docx")
    for p in paths:
        print(f"Wrote {p}")


if __name__ == "__main__":
    main()
