# -*- coding: utf-8 -*-
"""PPADUN9 분산 소스(web/android/ios) 상관관계 분석 .docx 생성"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

OUTPUT = Path(__file__).resolve().parent.parent / "PPADUN9_분산소스_상관관계.docx"


def set_cell_shading(cell, color_hex: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], "E8E8F0")
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return p


def build_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PPADUN9 분산 소스\n상관관계 분석 보고서")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1A, 0x6D, 0xFF)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("C:\\PPADUN9\\web  ·  android  ·  ios")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run("작성일: 2026년 6월 13일").font.size = Pt(11)

    doc.add_page_break()

    add_heading(doc, "1. 개요", 1)
    doc.add_paragraph(
        "PPADUN9(빠던9) 프로젝트 소스는 C:\\PPADUN9 아래 세 폴더로 분산되어 있습니다. "
        "역할은 서로 다르지만, 모두 같은 서비스(빠던9)를 구성합니다."
    )

    add_heading(doc, "1.1 폴더 구조", 2)
    add_code_block(
        doc,
        "C:\\PPADUN9\\\n"
        "├── web\\          ← 핵심: 웹 UI + 백엔드 API + DB 스키마 (본체)\n"
        "├── android\\      ← 유저 앱 Android 네이티브 껍데기 (APK 빌드용)\n"
        "└── ios\\          ← 유저 앱 iOS 네이티브 껍데기 (IPA 빌드용)\n"
        "\n"
        "(web 폴더 내부 추가 구성)\n"
        "web\\android-manager\\         ← 매니저 앱 Android\n"
        "web\\ios-manager-standalone\\  ← 매니저 앱 iOS",
    )

    add_table(
        doc,
        ["폴더", "역할", "대상 앱"],
        [
            ["web", "React 프론트 + Express 백엔드 + DB", "웹 + 서버 전체"],
            ["android", "Capacitor Android 프로젝트", "유저 앱 '빠던9'"],
            ["ios", "Capacitor iOS 프로젝트", "유저 앱 '빠던9'"],
            ["web/android-manager", "매니저용 Android", "매니저 앱"],
            ["web/ios-manager-standalone", "매니저용 iOS", "매니저 앱"],
        ],
    )

    add_heading(doc, "2. web 폴더 (본체)", 1)
    doc.add_paragraph(
        "web 폴더가 Source of Truth(진실의 원천)입니다. 모든 기능·로직·API가 여기서 개발됩니다."
    )
    add_table(
        doc,
        ["구성", "경로", "설명"],
        [
            ["프론트엔드", "client/", "User / Admin / Manager UI (React)"],
            ["백엔드", "server/", "REST API, WebSocket, Redis 세션"],
            ["DB 스키마", "shared/schema.ts", "Drizzle ORM + PostgreSQL"],
            ["배포", ".replit", "Replit Autoscale → ppamong.com"],
            ["GitHub", "hanbaedal/bbamong", "버전 관리 저장소"],
        ],
    )

    add_heading(doc, "3. android 폴더 (유저 앱 Android)", 1)
    add_table(
        doc,
        ["항목", "값"],
        [
            ["경로", "C:\\PPADUN9\\android"],
            ["App ID", "com.bbanden.nine"],
            ["앱 이름", "빠던9"],
            ["메인 클래스", "com.bbanden.nine.MainActivity"],
            ["딥링크", "ppadun9:// (소셜 로그인 콜백)"],
            ["서버 URL", "https://ppamong.com (Remote WebView)"],
        ],
    )
    doc.add_paragraph(
        "capacitor.config.json에 server.url이 설정되어 있어, 앱 실행 시 WebView가 "
        "프로덕션 서버(ppamong.com)를 로드합니다. web에서 서버만 배포해도 앱 UI가 자동 갱신됩니다."
    )
    doc.add_paragraph(
        "app/src/main/assets/public/ 안에는 예전 빌드 결과물(오프라인/폴백용)도 포함되어 있습니다."
    )

    add_heading(doc, "4. ios 폴더 (유저 앱 iOS)", 1)
    add_table(
        doc,
        ["항목", "값"],
        [
            ["경로", "C:\\PPADUN9\\ios"],
            ["Bundle ID", "com.bbanden.nine"],
            ["앱 이름", "빠던9"],
            ["내장 번들", "App/App/public/ (로컬 웹 파일)"],
            ["서버 URL", "capacitor.config.json에 server.url 없음"],
        ],
    )
    doc.add_paragraph(
        "Android와 달리 iOS는 server.url이 없어, App/App/public/에 번들된 웹 파일을 "
        "로컬로 실행하는 방식에 가깝습니다. web 빌드 후 iOS에 다시 sync가 필요할 수 있습니다."
    )

    add_heading(doc, "5. Android vs iOS 설정 차이", 2)
    add_table(
        doc,
        ["구분", "Android (C:\\PPADUN9\\android)", "iOS (C:\\PPADUN9\\ios)"],
        [
            ["JS 번들", "index-CZ8gc0IY.js", "index-Wy2pqOcY.js"],
            ["서버 연결", "https://ppamong.com 설정됨", "로컬 번들 위주"],
            ["UI 업데이트", "서버 배포만으로 반영 가능", "앱 재빌드 필요할 수 있음"],
        ],
    )

    add_heading(doc, "6. 실행·데이터 흐름", 1)
    add_code_block(
        doc,
        "[개발/수정]  C:\\PPADUN9\\web\n"
        "      ↓ npm run build / Replit deploy\n"
        "[프로덕션]  https://ppamong.com  ← Express + React + WebSocket\n"
        "      ↓ API / WS\n"
        "[DB]  Neon PostgreSQL  +  Redis\n"
        "\n"
        "[Android 앱]  WebView → ppamong.com (Remote URL)\n"
        "[iOS 앱]      로컬 public/ 또는 ppamong.com\n"
        "[웹 브라우저]  ppamong.com 직접 접속",
    )

    add_heading(doc, "7. 매니저 앱과의 관계", 1)
    add_table(
        doc,
        ["구분", "유저 앱", "매니저 앱"],
        [
            ["Android", "C:\\PPADUN9\\android", "web\\android-manager\\"],
            ["iOS", "C:\\PPADUN9\\ios", "web\\ios-manager-standalone\\"],
            ["App ID", "com.bbanden.nine", "com.ppadun9.manager"],
            ["접속 URL", "ppamong.com/login", "ppamong.com/manager"],
            ["딥링크", "ppadun9://", "ppadun9manager://"],
        ],
    )
    doc.add_paragraph(
        "유저 앱 Android/iOS만 web 밖(C:\\PPADUN9\\android, ios)으로 분리되어 있고, "
        "매니저 앱은 web 폴더 내부에 있습니다."
    )

    add_heading(doc, "8. 개발·빌드 순서", 1)

    add_heading(doc, "8.1 웹/API 수정", 2)
    add_code_block(
        doc,
        "cd C:\\PPADUN9\\web\n"
        "npm run dev          # 로컬 개발\n"
        "npm run build        # 프로덕션 빌드\n"
        "# → Replit 배포 → ppamong.com 반영",
    )

    add_heading(doc, "8.2 Android APK (유저 앱)", 2)
    add_code_block(
        doc,
        "cd C:\\PPADUN9\\web\n"
        "npm run build\n"
        "npx cap sync android   # (web에 android 연결 시)\n"
        "# Android Studio에서 C:\\PPADUN9\\android 열기 → APK 빌드",
    )

    add_heading(doc, "8.3 iOS IPA (유저 앱)", 2)
    add_code_block(
        doc,
        "cd C:\\PPADUN9\\web\n"
        "npm run build\n"
        "npx cap copy ios\n"
        "# Xcode에서 C:\\PPADUN9\\ios\\App 열기 → Archive",
    )

    add_heading(doc, "9. 분산 구조 주의사항", 1)
    add_table(
        doc,
        ["이슈", "설명"],
        [
            ["소스 중복", "android/ios의 assets/public은 web 빌드 결과의 복사본"],
            ["버전 불일치", "Android JS와 iOS JS 번들 해시가 서로 다름"],
            ["설정 불일치", "Android는 Remote URL, iOS는 로컬 번들 혼용"],
            ["Git 범위", "GitHub bbamong에는 web만 push됨. android/ios는 별도 관리"],
            ["매니저 위치", "매니저는 web 내부, 유저 앱은 web 외부"],
        ],
    )

    add_heading(doc, "10. 요약", 1)
    add_table(
        doc,
        ["질문", "답"],
        [
            ["세 폴더가 같은 프로그램인가?", "예. 같은 빠던9 서비스"],
            ["어디가 본체인가?", "web (UI + API + DB)"],
            ["android/ios 역할?", "스토어 배포용 Capacitor 네이티브 껍데기"],
            ["서버 연결?", "모두 https://ppamong.com"],
            ["GitHub?", "web만 hanbaedal/bbamong에 등록"],
        ],
    )

    doc.add_paragraph()
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = summary.add_run(
        "한 줄 요약: web = 두뇌·몸통, android/ios = 스마트폰 앱 포장지. "
        "기능 수정은 web에서, 스토어 출시는 android/ios에서 APK/IPA 빌드."
    )
    sr.italic = True
    sr.font.size = Pt(10)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("— 보고서 끝 —")
    fr.font.size = Pt(10)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    return doc


def main():
    doc = build_document()
    doc.save(str(OUTPUT))
    print(f"생성 완료: {OUTPUT}")


if __name__ == "__main__":
    main()
