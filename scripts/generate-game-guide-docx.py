# -*- coding: utf-8 -*-
"""PPAMONG 야구 예측 게임 상세 설명 .docx 생성"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent.parent / "docs" / "PPAMONG_게임_방법.docx"


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    doc = Document()

    title = doc.add_heading("PPAMONG 야구 예측 게임 — 상세 설명", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "PPAMONG(빠몽)의 핵심은 실시간 야구 경기 중 「다음 타석 결과」를 맞추는 라운드형 예측 게임입니다. "
        "회원은 포인트(앱에서는 「참여 기록」)를 걸고 예측하고, 맞추면 다른 참가자들이 건 포인트를 나눠 받습니다."
    )

    # 1
    doc.add_heading("1. 게임의 기본 개념", level=1)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["예측 대상", "다음 타자의 타석 결과"],
            ["선택지", "1루 · 2루 · 3루 · 홈런 · 아웃 (5가지)"],
            ["단위", "라운드 — 매니저가 한 번씩 열고 닫는 예측 구간"],
            ["베팅", "라운드당 기본 100P"],
            ["운영", "매니저가 라운드 시작·중지·결과 입력, 관리자가 경기·광고·회원 전체 관리"],
        ],
    )
    doc.add_paragraph(
        "경기 전체가 아니라 「지금 이 타석은 어떻게 끝날까?」를 맞추는 구조입니다. "
        "라운드는 이닝과 1:1이 아니라, 타석(또는 매니저가 정한 예측 단위)마다 진행됩니다."
    )

    # 2
    doc.add_heading("2. 참여 전 준비 (회원)", level=1)
    doc.add_heading("로그인", level=2)
    add_bullets(
        doc,
        [
            "일반 로그인, 카카오·구글·애플 소셜 로그인, 게스트 로그인 지원",
            "게스트는 일부 기능(출석 등) 제한 가능",
        ],
    )
    doc.add_heading("참여 기록(포인트) 확보", level=2)
    doc.add_paragraph("예측에 쓰려면 최소 100P가 필요합니다.")
    add_table(
        doc,
        ["방법", "지급량"],
        [
            ["회원가입", "+1,000P"],
            ["유효한 추천인 코드 입력", "+1,000P (추가)"],
            ["일일 출석 (/attendance)", "+100P (하루 1회)"],
            ["소개 영상 시청", "1번: 500P, 2번: 1,000P, 3번: 1,500P (준비 중)"],
            ["예측 승리", "라운드별 정산 (아래 참고)"],
        ],
    )
    doc.add_paragraph("포인트 내역은 /point-history, 승리 기록은 설정 → 승리 내역에서 확인할 수 있습니다.")

    # 3
    doc.add_heading("3. 회원 플레이 흐름 (/prediction)", level=1)

    doc.add_heading("① 경기 선택", level=2)
    add_bullets(
        doc,
        [
            "오늘 진행 중이거나 예정된 경기 목록 표시",
            "종료된 경기(completed)는 선택 불가",
            "경기 선택 후 WebSocket으로 해당 경기 실시간 이벤트 수신",
        ],
    )

    doc.add_heading("② 예측 대기", level=2)
    add_bullets(
        doc,
        [
            "매니저가 「예측 시작」하기 전에는 선택 버튼이 비활성",
            "화면: 「다음 타자 예측을 기다리는 중 입니다.」",
            "predictionEnabled: false 상태",
        ],
    )

    doc.add_heading("③ 예측 선택 및 베팅", level=2)
    doc.add_paragraph("매니저가 예측을 열면(prediction_started):")
    add_bullets(
        doc,
        [
            "1루 / 2루 / 3루 / 홈런 / 아웃 중 하나 선택",
            "확인 팝업에서 경기장·일시·예측·100P·잔여 기록 확인",
            "확정 시 즉시 100P 차감 (첫 제출 시에만)",
        ],
    )
    doc.add_paragraph("중요 규칙:")
    add_bullets(
        doc,
        [
            "같은 라운드에서 선택 변경 → 추가 차감 없음 (이미 건 100P 유지)",
            "예측 취소 → 현재 라운드·결과 확정 전이면 100P 전액 환불",
            "결과 확정 후 → 변경·취소 불가",
            "포인트 부족 시 「참여기회가 부족합니다」 오류",
        ],
    )

    doc.add_heading("④ 결과 대기", level=2)
    add_bullets(
        doc,
        [
            "예측 제출 후 매니저가 예측 중지 → 결과 입력할 때까지 대기",
            "WebSocket round_result로 개인별 승패·획득 포인트 수신",
            "연결 끊김 시 폴링으로 보완 (최대 약 7분)",
        ],
    )

    doc.add_heading("⑤ 결과 화면", level=2)
    add_table(
        doc,
        ["결과", "화면", "후속"],
        [
            ["성공", "획득 포인트 표시", "5초 후 자동 종료, 기부 선택 가능"],
            ["실패", "실패 안내", "3초 후 자동 종료"],
            ["예측 중지", "「예측이 종료되었습니다」", "다음 라운드 대기"],
        ],
    )

    doc.add_heading("⑥ 기부 (선택)", level=2)
    doc.add_paragraph("승리 시 상금의 10%를 기부할 수 있습니다.")
    add_bullets(
        doc,
        [
            "기부하지 않기 / 10% 기부 중 선택",
            "기부 대상 = (획득액 - 원금 100P)의 10% (반올림)",
            "예: 상금 200P → 기부 20P",
            "라운드당 1회만 기부 가능",
            "기부 내역·랭킹: 설정 → 기부 내역",
        ],
    )

    # 4
    doc.add_heading("4. 포인트 정산 (핵심 로직)", level=1)
    doc.add_heading("패배자", level=2)
    doc.add_paragraph("베팅한 100P는 이미 차감된 상태 → 추가 차감 없음")

    doc.add_heading("승리자", level=2)
    p = doc.add_paragraph()
    p.add_run("패자 풀 = 틀린 사람들이 건 포인트 합계\n").bold = True
    p.add_run("1인당 상금 = 패자 풀 ÷ 승자 수\n")
    p.add_run("상금(반올림) = 10P 단위로 올림 (ceil)\n")
    p.add_run("최종 지급 = 원금 100P + 상금")

    doc.add_heading("정산 예시", level=2)
    doc.add_paragraph("예시 1 — 참가 10명, 각 100P, 승자 2명:")
    add_bullets(
        doc,
        [
            "패자 풀 = 800P",
            "1인 상금 = 800 ÷ 2 = 400P",
            "지급 = 100 + 400 = 500P",
        ],
    )
    doc.add_paragraph("예시 2 — 승자 1명:")
    add_bullets(doc, ["패자 9명 → 900P 전부 1명에게 (+ 원금 100P)"])
    doc.add_paragraph("예시 3 — 승자 0명:")
    add_bullets(doc, ["아무도 맞추지 못하면 지급 없음, 패자들의 100P는 소멸"])
    doc.add_paragraph("상금은 10P 단위 올림 처리됩니다.")

    # 5
    doc.add_heading("5. 라운드 생명주기 (매니저 운영)", level=1)
    doc.add_paragraph("매니저 앱(/manager/match/:id)에서 경기를 진행합니다.")
    flow = doc.add_paragraph()
    flow.add_run(
        "경기 시작\n"
        "    ↓\n"
        "[라운드 N]\n"
        "  ① 예측 시작  → predictionEnabled = true, 회원 예측 가능\n"
        "  ② 예측 중지  → predictionEnabled = false, 더 이상 신규·변경 불가\n"
        "  ③ 결과 입력  → 5가지 중 실제 타석 결과 선택\n"
        "  ④ 자동 정산  → 승패·포인트 지급, round_next로 라운드 N+1\n"
        "    ↓\n"
        "(반복)\n"
        "    ↓\n"
        "경기 종료"
    ).font.name = "Consolas"

    doc.add_heading("매니저 버튼별 의미", level=2)
    add_table(
        doc,
        ["동작", "효과"],
        [
            ["예측 시작", "해당 라운드 예측 오픈, 진행 중 광고 자동 중지"],
            ["예측 중지", "예측 마감 (결과 입력 전 단계)"],
            ["결과 전송", "정산 + 다음 라운드로 자동 이동"],
            ["다음 라운드 (강제)", "예측 중이어도 중지 후 라운드 증가 가능"],
        ],
    )

    doc.add_heading("예외 처리", level=2)
    add_bullets(
        doc,
        [
            "예측 재시작 (중지 후 다시 시작): 기존 pending 예측 전원 환불 후 초기화",
            "결과 없이 다음 라운드: 강제 진행 시 이전 라운드 결과 생략 가능",
            "경기 취소·종료: 신규 예측 불가",
        ],
    )

    # 6
    doc.add_heading("6. 실시간 연동 (WebSocket)", level=1)
    doc.add_paragraph("회원·매니저 앱은 경기별 WebSocket(SSE 기반 broadcast)으로 동기화됩니다.")
    add_table(
        doc,
        ["이벤트", "의미"],
        [
            ["prediction_started", "예측 시작, 선택 가능"],
            ["prediction_stopped", "예측 마감"],
            ["round_result", "라운드 결과·개인별 wonAmount"],
            ["round_next", "다음 라운드로 이동"],
            ["match_end", "경기 종료"],
            ["ad_started / ad_stopped", "라운드 사이 광고 재생"],
            ["stats_update", "참가자·베팅 통계 갱신"],
        ],
    )
    doc.add_paragraph("관리자는 실시간 게임 모니터링 화면에서 전체 현황을 볼 수 있습니다.")

    # 7
    doc.add_heading("7. 앱 구조와 화면 구분", level=1)
    add_table(
        doc,
        ["화면", "경로", "역할"],
        [
            ["경기 참여", "/prediction", "예측 게임 본체"],
            ["게임 가이드", "/home/game-guide", "관리자 등록 설명 + 참여 버튼"],
            ["빠몽이의 보물창고", "/, /home", "쇼핑·소개 (게임과 분리)"],
        ],
    )
    add_bullets(
        doc,
        [
            "상단 빠몽이 로고: 게임 ↔ 보물창고 이동 (로그인 상태 유지)",
            "게임 화면: 하단 메뉴(초대·출석·게시·추가·로그아웃)",
            "보물창고: 게임용 하단 메뉴 없음",
        ],
    )

    # 8
    doc.add_heading("8. 역할별 책임", level=1)
    doc.add_heading("회원 (User 앱 / APK)", level=2)
    add_bullets(
        doc,
        [
            "경기 선택 → 예측 → 결과 확인 → (선택) 기부",
            "출석·영상·초대로 포인트 확보",
            "커뮤니티 게시판, 고객센터, 공지 등 부가 기능",
        ],
    )
    doc.add_heading("매니저 (Manager 앱 / APK)", level=2)
    add_bullets(
        doc,
        [
            "배정된 경기만 운영",
            "라운드 시작·중지·결과 입력",
            "실시간 참가 현황 확인",
            "관리자 승인 후 사용 가능",
        ],
    )
    doc.add_heading("관리자 (Admin 웹)", level=2)
    add_bullets(
        doc,
        [
            "경기·구장·회원·포인트·광고 등록",
            "매니저 승인·경기 배정",
            "홈페이지 게임 가이드 문구·이미지 설정",
            "라운드 사이 광고 송출 제어",
        ],
    )

    # 9
    doc.add_heading("9. 자주 묻는 상황", level=1)
    faqs = [
        ("예측을 안 했는데 포인트가 줄었나요?", "예측 확정 시점에만 100P가 차감됩니다. 대기만 하면 차감 없습니다."),
        ("선택을 바꿨는데 200P가 빠졌나요?", "같은 라운드 변경은 추가 차감 없음. 내역을 확인해 보세요."),
        ("예측했는데 결과가 안 나와요", "매니저가 아직 예측 중지·결과 입력을 하지 않았을 수 있습니다. 네트워크 끊김 시 앱이 폴링으로 재시도합니다."),
        ("맞췄는데 기부 팝업이 안 뜨나요?", "획득액이 원금(100P) 이하이면 기부 대상 상금이 없어 팝업이 생략될 수 있습니다."),
        ("한 경기에 여러 번 참여할 수 있나요?", "라운드마다 1회입니다. 경기가 진행되는 동안 라운드가 반복됩니다."),
    ]
    for q, a in faqs:
        pq = doc.add_paragraph()
        pq.add_run(f"Q. {q}").bold = True
        doc.add_paragraph(f"A. {a}")

    # 10
    doc.add_heading("10. 한 줄 요약", level=1)
    summary = doc.add_paragraph()
    run = summary.add_run(
        "PPAMONG은 매니저가 실시간으로 여는 타석 단위 예측 라운드에 회원이 100P를 걸고 "
        "1루·2루·3루·홈런·아웃 중 하나를 고르는 게임입니다. "
        "맞추면 틀린 사람들의 포인트를 나눠 받고, 틀리면 베팅한 포인트를 잃습니다. "
        "WebSocket으로 즉시 동기화되며, 승리 후에는 선택적으로 상금의 10%를 기부할 수 있습니다."
    )
    run.italic = True

    # footer note
    doc.add_paragraph()
    note = doc.add_paragraph("문서 버전: 2026-06-17 | PPAMONG (ppamong.com)")
    note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in note.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
