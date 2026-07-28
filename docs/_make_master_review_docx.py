# -*- coding: utf-8 -*-
"""PPAMONG 전체 점검 마스터 DOCX 생성 (2026-07 기준)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=18 if level == 1 else 14 if level == 2 else 12, bold=True)
    return p


def add_para(doc, text, size=11, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run, size=11)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=9)
    doc.add_paragraph()


def add_flow_box(doc, title, steps):
    add_para(doc, title, bold=True, size=11)
    add_numbered(doc, steps)
    doc.add_paragraph()


def add_screenshot_placeholder(doc, label, description):
    add_para(doc, f"▶ {label}", bold=True)
    add_para(doc, description, size=10)
    p = doc.add_paragraph()
    run = p.add_run("[ 화면 캡처 삽입 위치 — ppamong.com 실제 스크린샷 ]")
    set_run_font(run, size=10, color=RGBColor(0x99, 0x99, 0x99))
    p.paragraph_format.left_indent = Inches(0.25)
    doc.add_paragraph()


def add_toc_manual(doc):
    add_heading(doc, "목차", 1)
    items = [
        "1. 한눈에 보기 (Executive Summary)",
        "2. 시스템 구조",
        "3. 예측 게임 흐름 (관리자 · 운영자 · 사용자)",
        "4. gamePhase · 운영자 3버튼",
        "5. 데이터베이스 (MongoDB)",
        "6. API · WebSocket 요약",
        "7. 화면 흐름 (캡처 가이드)",
        "8. 제약 · 점검 체크리스트",
        "부록 A. 전체 컬렉션 목록",
        "부록 B. API 목록 (확장)",
        "부록 C. 배포 · Secrets",
    ]
    add_bullets(doc, items)
    doc.add_page_break()


def main():
    out = Path(__file__).resolve().parent / "PPAMONG_전체점검_마스터.docx"
    doc = Document()

    # 표지
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("빠몽이 (PPAMONG)")
    set_run_font(run, size=24, bold=True, color=RGBColor(0xE1, 0x19, 0x36))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("전체 점검 · 시스템 · DB · 게임 흐름\n마스터 설명서")
    set_run_font(run, size=16, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "기준일: 2026-07-26\n"
        "GitHub: hanbaedal/bbamong (main)\n"
        "호스팅: Replit Autoscale · DB: MongoDB Atlas · 세션: Redis\n"
        "※ db-erd.md(PostgreSQL)는 레거시 참고용 — 운영 DB는 MongoDB"
    )
    set_run_font(run, size=10, color=RGBColor(0x66, 0x66, 0x66))
    doc.add_page_break()

    add_toc_manual(doc)

    # ── 1. 한눈에 보기 ──
    add_heading(doc, "1. 한눈에 보기 (Executive Summary)", 1)
    add_para(
        doc,
        "PPAMONG은 KBO 실시간 야구 타석 예측 게임 + 포인트·출석·게시판 + 쇼핑몰을 "
        "하나의 Express 서버(ppamong.com)에서 제공합니다. 타석 결과는 API로 제공되지 않아 "
        "운영자(매니저)의 수동 조작이 필수이며, API-SPORTS는 일정·점수·경기 종료만 자동화합니다.",
    )
    add_table(
        doc,
        ["역할", "클라이언트", "핵심 URL", "하는 일"],
        [
            ["일반 사용자", "Android/iOS 앱", "/prediction, /home", "타석·사이드 배팅, 포인트, 광고"],
            ["운영자", "매니저 앱", "/manager/*", "예측 시작/중지/결과, 3버튼 라운드 진행"],
            ["관리자", "웹", "/admin/*", "경기 달력, 회원·몰, 모니터링"],
            ["쇼핑", "웹", "/shop", "굿즈 주문 (정회원)"],
        ],
    )
    add_table(
        doc,
        ["자동 (API/시스템)", "수동 (운영자/관리자)"],
        [
            ["KBO 일정·점수·종료", "타석 결과(1루~아웃)"],
            ["DB 캐시 우선 일정 적재", "예측 시작·중지·결과 입력"],
            ["사이드배팅 1회 시작 마감", "gamePhase 3버튼(다음타자/투수교체/공수교대)"],
            ["경기 종료 시 사이드 정산", "비상 시 controlMode=manual"],
        ],
    )

    # ── 2. 시스템 구조 ──
    add_heading(doc, "2. 시스템 구조", 1)
    add_para(doc, "클라이언트 4종 → ppamong.com (Node.js + Express) → MongoDB / Redis / API-SPORTS", bold=True)
    add_bullets(
        doc,
        [
            "단일 서버: UserApp · ManagerApp · AdminApp · MallApp 모두 동일 백엔드",
            "실시간: WebSocket /ws/match (JWT) + API-SPORTS 폴링 약 2.5초",
            "인증: JWT httpOnly 쿠키, OAuth(카카오/구글/애플), 운영자 일회용 로그인 링크",
            "광고: AdMob (공수교대 보상·타자교체 배너·전면)",
            "코드: client/ (React) · server/ (Express) · shared/ (타입·배당)",
        ],
    )
    add_heading(doc, "2.1 외부 연동", 2)
    add_table(
        doc,
        ["시스템", "용도", "비고"],
        [
            ["API-SPORTS Baseball", "KBO 일정·이닝·점수·FT", "초/말(공수) 미제공"],
            ["MongoDB Atlas", "회원·경기·배팅·몰", "운영 DB"],
            ["Redis", "세션·소셜 pending", "Replit 내장 가능"],
            ["AdMob", "앱 광고", "보상 500P 등"],
        ],
    )

    # ── 3. 게임 흐름 ──
    add_heading(doc, "3. 예측 게임 흐름", 1)

    add_heading(doc, "3.1 관리자 — 경기 일정", 2)
    add_flow_box(
        doc,
        "달력 → DB 캐시 우선 동기화",
        [
            "/admin/match-management 달력에서 날짜 클릭",
            "POST /api/admin/matches/sync-from-api-sports { date }",
            "해당일 ApiSportsScheduleCache(DB) 우선 → 없으면 API 호출 후 캐시 저장",
            "Match 1~5경기 생성/갱신, apiSportsGameId·팀명·구장 연결",
            "운영자 op1~op5 registrationOrder 자동 배정",
            "모달에서 경기 목록·모니터링 링크 확인",
        ],
    )

    add_heading(doc, "3.2 운영자 — 타석 라운드 루프", 2)
    add_flow_box(
        doc,
        "한 타석(라운드) 사이클",
        [
            "매니저 앱 로그인 (링크 1회용 또는 ID/PW)",
            "배정 경기 입장, API 스코어보드 확인",
            "「예측 시작」→ predictionEnabled=true, WS prediction_started",
            "사용자 배팅 (50~1000P, 1루/2루/3루/홈런/아웃)",
            "「예측 중지」→ 배팅 마감",
            "「결과 입력」(1루~아웃) → 정산, WS round_result",
            "다음 타자 / 투수 교체 / 공수 교대 중 선택 → WS round_next",
            "경기 종료: API FT 또는 관리자 종료 → 사이드배팅 정산",
        ],
    )

    add_heading(doc, "3.3 사용자 — 참여 흐름", 2)
    add_flow_box(
        doc,
        "앱 예측 참여",
        [
            "로그인 → 홈 → 「경기 참여」→ /prediction",
            "(선택) 승리팀·최종스코어 사이드 배팅 — 1회 시작 전까지",
            "경기 선택 → 예측 옵션·금액 → 확인 → POST predictions",
            "대기 화면: gamePhase(회·공격팀·N번째 타자) + 빠몽이 캐릭터",
            "결과: 성공/실패 → (선택) 기부 → 다음 라운드 대기",
            "공수교대: 보상 광고(500P), 타자교체: 하단 배너",
        ],
    )

    add_heading(doc, "3.4 배당 · 사이드 배팅", 2)
    add_table(
        doc,
        ["구분", "선택", "배당", "비고"],
        [
            ["타석", "아웃", "1.2×", "50~1000P"],
            ["타석", "1루", "1.5×", ""],
            ["타석", "2루", "3×", ""],
            ["타석", "3루", "10×", ""],
            ["타석", "홈런", "5×", ""],
            ["사이드", "승리팀", "2×", "100P 단위, 1회 시작 마감"],
            ["사이드", "최종 스코어", "20×", "정확 일치 시"],
        ],
    )

    # ── 4. gamePhase ──
    add_heading(doc, "4. gamePhase · 운영자 3버튼", 1)
    add_para(
        doc,
        "Match 컬렉션의 gameInning, inningHalf(top/bottom), batterIndexInHalf로 "
        "회원 대기 화면에 「N경기 (M회) / 원정팀 공격 (K번째 타자)」 형태로 표시합니다. "
        "API-SPORTS와 별개이며 운영자 버튼으로만 갱신됩니다.",
    )
    add_table(
        doc,
        ["버튼", "API", "gamePhase 변화", "용도"],
        [
            ["다음 타자", "POST .../round/next-batter", "batterIndexInHalf +1", "타순 진행"],
            ["투수 교체", "POST .../round/pitcher-change", "변화 없음", "같은 타자·라운드만"],
            ["공수 교대", "POST .../round/switch-half", "초↔말, 타순 1", "이닝 half 전환"],
        ],
    )
    add_bullets(
        doc,
        [
            "공수교대 API 자동 연동 없음 (API에 초/말 정보 없음)",
            "투수 교체는 타순·이닝을 바꾸지 않음 — 라운드·배너만 진행",
        ],
    )

    # ── 5. DB ──
    add_heading(doc, "5. 데이터베이스 (MongoDB Atlas)", 1)
    add_para(doc, "스키마: server/mongodb/models.ts · DB명: MONGODB_DB_NAME (기본 ppamong)")

    add_heading(doc, "5.1 게임 핵심 컬렉션 (상세)", 2)

    add_para(doc, "Match — 경기", bold=True)
    add_table(
        doc,
        ["필드", "타입/값", "설명"],
        [
            ["id", "string UUID", "경기 PK"],
            ["name", "string", "1경기~5경기"],
            ["matchDate", "string YYYY-MM-DD", "KST 경기일"],
            ["stadiumId", "number", "→ Stadium.id"],
            ["matchStatus", "scheduled|ongoing|completed|cancelled", ""],
            ["currentRound", "number", "현재 라운드 번호"],
            ["predictionEnabled", "boolean", "예측 접수 중"],
            ["registrationOrder", "1~5", "운영자 op 배정 순서"],
            ["apiSportsGameId", "number|null", "API 연결 ID"],
            ["liveScoreboard", "object", "R-H-E, inning, status"],
            ["controlMode", "auto|manual", "API 자동 vs 수동"],
            ["sideBetsLocked", "boolean", "1회 시작 후 사이드 마감"],
            ["gameInning", "number", "운영자 기준 회"],
            ["inningHalf", "top|bottom", "초/말"],
            ["batterIndexInHalf", "number", "공수 내 타순"],
        ],
    )

    add_para(doc, "Prediction — 타석 배팅", bold=True)
    add_table(
        doc,
        ["필드", "설명"],
        [
            ["userId + matchId + roundNumber", "유니크 (라운드당 1회)"],
            ["prediction", "1루|2루|3루|홈런|아웃"],
            ["amount", "배팅 포인트"],
            ["status", "pending → success|fail"],
            ["wonAmount", "적중 지급"],
        ],
    )

    add_para(doc, "MatchSideBet — 승리팀/스코어", bold=True)
    add_table(
        doc,
        ["필드", "설명"],
        [
            ["type", "winner | score"],
            ["winnerPick", "home | away"],
            ["homeScorePick, awayScorePick", "스코어 배팅"],
            ["odds", "2 또는 20"],
            ["status", "pending → success|fail|refunded"],
        ],
    )

    add_para(doc, "ApiSportsScheduleCache — 일정 DB 캐시", bold=True)
    add_table(
        doc,
        ["필드", "설명"],
        [
            ["matchDate + apiSportsGameId", "유니크"],
            ["homeTeamName, awayTeamName", "팀명"],
            ["statusShort, statusLong", "API 상태"],
            ["homeScore, awayScore", "점수"],
            ["fetchedAt", "적재 시각"],
        ],
    )

    add_para(doc, "RoundStatistics — 라운드 통계", bold=True)
    add_bullets(
        doc,
        [
            "matchId + roundNumber별 참여자·포인트·당첨 수",
            "predictionStartTime / predictionStopTime",
            "isPredictionStarted / isPredictionStopped / isResultSent",
        ],
    )

    add_heading(doc, "5.2 회원 · 운영 (요약)", 2)
    add_table(
        doc,
        ["컬렉션", "용도"],
        [
            ["User", "회원, points, OAuth, inviteCode"],
            ["AdminUser", "관리자·슈퍼바이저"],
            ["Manager (operators)", "op1~op5, apiSyncEnabled, loginLink"],
            ["Stadium", "구장명"],
            ["PointTransaction", "포인트 이력"],
            ["HomePageSettings", "앱 홈·게임 설명 문구"],
        ],
    )

    add_heading(doc, "5.3 쇼핑몰 (요약)", 2)
    add_table(
        doc,
        ["컬렉션", "용도"],
        [
            ["GoodsCategory, GoodsProduct", "상품·카테고리"],
            ["MallOrder", "주문"],
            ["MallWarehouse, MallLocation, MallStockMovement", "재고"],
            ["MallSupplier, MallPurchaseOrder", "구매·발주"],
            ["ShopInquiry", "쇼핑몰 문의"],
        ],
    )

    # ── 6. API ──
    add_heading(doc, "6. API · WebSocket 요약", 1)
    add_para(doc, "전체 명세: docs/api-spec.md (별도 800줄+)")

    add_heading(doc, "6.1 게임 · 경기 (핵심)", 2)
    add_table(
        doc,
        ["메서드", "경로", "설명"],
        [
            ["POST", "/api/admin/matches/sync-from-api-sports", "날짜 일정 DB·경기 연결"],
            ["POST", "/api/admin/matches/import-season-schedule", "시즌 일정 DB 적재"],
            ["GET", "/api/admin/matches", "관리자 경기 목록"],
            ["GET", "/api/matches/:id", "경기 상세·gamePhase"],
            ["GET", "/api/matches/:id/scoreboard", "라이브 스코어"],
            ["GET", "/api/api-sports/health", "API 헬스"],
            ["POST", "/api/live-match/predictions", "타석 배팅"],
            ["GET", "/api/live-match/predictions/:matchId/check", "내 예측 상태"],
            ["POST", "/api/live-match/side-bets", "사이드 배팅"],
        ],
    )

    add_heading(doc, "6.2 운영자 (매니저)", 2)
    add_table(
        doc,
        ["메서드", "경로", "설명"],
        [
            ["POST", "/api/manager/login-with-link", "일회용 링크 로그인"],
            ["GET", "/api/manager/matches/today", "오늘 배정 경기"],
            ["POST", "/api/manager/matches/:id/prediction/start", "예측 시작"],
            ["POST", "/api/manager/matches/:id/prediction/stop", "예측 중지"],
            ["POST", "/api/manager/matches/:id/result", "결과 입력"],
            ["POST", "/api/manager/control/:id/round/next-batter", "다음 타자"],
            ["POST", "/api/manager/control/:id/round/pitcher-change", "투수 교체"],
            ["POST", "/api/manager/control/:id/round/switch-half", "공수 교대"],
        ],
    )

    add_heading(doc, "6.3 WebSocket (/ws/match)", 2)
    add_table(
        doc,
        ["이벤트", "방향", "설명"],
        [
            ["prediction_started", "S→C", "예측 시작"],
            ["prediction_stopped", "S→C", "예측 중지"],
            ["round_result", "S→C", "결과·정산"],
            ["round_next", "S→C", "라운드 전환 + gamePhase"],
            ["banner_ad_show", "S→C", "배너 광고"],
            ["rewarded_ad_offer", "S→C", "보상 광고 제안"],
            ["scoreboard_update", "S→C", "스코어 갱신"],
            ["match_ended", "S→C", "경기 종료"],
        ],
    )

    # ── 7. 화면 캡처 가이드 ──
    add_heading(doc, "7. 화면 흐름 (캡처 가이드)", 1)
    add_para(
        doc,
        "아래는 실제 ppamong.com / 앱에서 캡처하여 삽입할 화면 목록입니다. "
        "캡처 파일을 docs/screenshots/ 에 저장 후 Word에서 「그림 삽입」으로 교체하세요.",
    )

    add_heading(doc, "7.1 사용자 앱", 2)
    add_screenshot_placeholder(doc, "U-01 홈", "인사·경기 참여 버튼·사용설명서")
    add_screenshot_placeholder(doc, "U-02 경기 선택", "/prediction 경기 리스트")
    add_screenshot_placeholder(doc, "U-03 배팅", "1루~아웃 선택·금액·확인")
    add_screenshot_placeholder(doc, "U-04 대기", "gamePhase 2줄 + 빠몽이 + 「진루 예측을 기다리고 있습니다」")
    add_screenshot_placeholder(doc, "U-05 결과", "성공/실패·기부 선택")

    add_heading(doc, "7.2 운영자 앱", 2)
    add_screenshot_placeholder(doc, "M-01 로그인", "링크 로그인·음성 안내")
    add_screenshot_placeholder(doc, "M-02 경기 상세", "스코어보드·예측 시작/중지/결과")
    add_screenshot_placeholder(doc, "M-03 3버튼", "다음 타자 · 투수 교체 · 공수 교대")

    add_heading(doc, "7.3 관리자 웹", 2)
    add_screenshot_placeholder(doc, "A-01 사이트맵", "헤더 사이트맵 · 5열 트리")
    add_screenshot_placeholder(doc, "A-02 경기 달력", "/admin/match-management 한 화면 달력")
    add_screenshot_placeholder(doc, "A-03 일정 모달", "날짜 클릭 후 경기 표")

    # ── 8. 제약 · 체크리스트 ──
    add_heading(doc, "8. 제약 · 점검 체크리스트", 1)
    add_table(
        doc,
        ["등급", "항목", "대응"],
        [
            ["높음", "타석 결과 API 없음", "운영자 수동 결과"],
            ["높음", "API 미연결", "달력 sync·팀명 확인"],
            ["높음", "Deploy ≠ Run", "Publishing Redeploy 필수"],
            ["중간", "공수 API 자동 불가", "운영자 공수교대 버튼"],
            ["중간", "미정산 라운드", "종료 전 결과 입력"],
            ["낮음", "Free API 플랜", "Pro 키 권장"],
        ],
    )
    add_heading(doc, "8.1 당일 운영 체크", 2)
    add_bullets(
        doc,
        [
            "Replit Secrets: MONGODB_URI, JWT_*, API_SPORTS_KEY",
            "git fetch && reset --hard origin/main → Redeploy",
            "GET /api/api-sports/health → healthy",
            "달력에서 오늘 sync → 1~5경기·API 연결",
            "운영자 op1~op2 API sync ON 확인",
            "운영자 3버튼·gamePhase 회원 화면 확인",
        ],
    )

    doc.add_page_break()

    # ── 부록 A ──
    add_heading(doc, "부록 A. 전체 MongoDB 컬렉션", 1)
    collections = [
        "User", "Stadium", "Match", "ApiSportsScheduleCache",
        "Prediction", "MatchSideBet", "RoundStatistics",
        "AttendanceRecord", "Post", "Comment", "PointTransaction",
        "Inquiry", "Notice", "Term", "Faq", "Ebook", "EbookPurchase",
        "AdminUser", "WaitingScreen", "Advertisement", "AdViewHistory",
        "HomePageSettings", "GoodsCategory", "GoodsProduct",
        "AppAdmobConfig", "ShopInquiry", "MallProductReview",
        "MallOrder", "MallWarehouse", "MallLocation", "MallStockMovement",
        "MallSupplier", "MallPurchaseOrder", "Counter",
    ]
    add_bullets(doc, collections)

    # ── 부록 B ──
    add_heading(doc, "부록 B. API 목록 (확장 참고)", 1)
    add_para(doc, "상세 Request/Response는 docs/api-spec.md 참조.")
    add_bullets(
        doc,
        [
            "인증: /api/login, /api/signup, /api/auth/*/callback",
            "유저: /api/users/me, 출석, 포인트 이력",
            "게시판·문의·공지·약관·FAQ·이북",
            "관리자: /api/admin/* (회원·직원·경기·광고·홈설정)",
            "몰: /api/shop/*, /api/admin/mall/*",
            "매니저: /api/manager/*",
        ],
    )

    # ── 부록 C ──
    add_heading(doc, "부록 C. 배포 · Secrets", 1)
    add_table(
        doc,
        ["Secret", "필수", "설명"],
        [
            ["MONGODB_URI", "○", "Atlas 연결 문자열"],
            ["JWT_SECRET / JWT_REFRESH_SECRET", "○", "서로 다른 값"],
            ["BASE_URL", "○", "https://ppamong.com"],
            ["API_SPORTS_KEY", "○", "KBO 일정·스코어"],
            ["API_SPORTS_KBO_LEAGUE_ID", "△", "기본 5"],
            ["KAKAO/GOOGLE OAuth", "△", "소셜 로그인"],
        ],
    )
    add_bullets(
        doc,
        [
            "배포: git fetch origin main && git reset --hard origin/main",
            "Publishing → Autoscale → Redeploy (5~15분 소요 가능)",
            "상세: docs/PPAMONG_DEPLOY_CHECKLIST.md",
        ],
    )

    footer = doc.add_paragraph()
    run = footer.add_run("— 문서 끝 —  PPAMONG 전체점검 마스터 v2026-07-26")
    set_run_font(run, size=10, color=RGBColor(0x88, 0x88, 0x88))
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
