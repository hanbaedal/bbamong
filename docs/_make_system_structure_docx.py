# -*- coding: utf-8 -*-
"""빠몽이(ppamong) 시스템 구조·설명 DOCX 생성."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13 if level == 2 else 12, bold=True)
    return p


def add_para(doc, text, size=11, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
    doc.add_paragraph()


def main():
    out = Path(__file__).resolve().parent / "PPAMONG_시스템구조_설명.docx"
    doc = Document()

    # 표지
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("빠몽이 (ppamong.com)")
    set_run_font(run, size=22, bold=True, color=RGBColor(0x20, 0x1E, 0x22))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("시스템 구조 · 기능 설명서")
    set_run_font(run, size=16, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "기준일: 2026-07-25\n"
        "GitHub: hanbaedal/bbamong (main)\n"
        "호스팅: Replit Autoscale · DB: MongoDB Atlas"
    )
    set_run_font(run, size=10, color=RGBColor(0x66, 0x66, 0x66))

    doc.add_paragraph()

    # 1. 개요
    add_heading(doc, "1. 제품 개요", 1)
    add_para(
        doc,
        "빠몽이는 실시간 야구(KBO) 타석 예측·경기 단위 사이드 배팅과 포인트·출석·게시판·쇼핑몰을 "
        "하나의 백엔드에서 제공하는 서비스입니다. 백엔드 API는 Replit의 web 서버 하나에서 처리하며, "
        "클라이언트는 사용자 앱·운영자(매니저) 앱·관리자 웹·쇼핑몰 웹으로 구성됩니다.",
    )
    add_table(
        doc,
        ["구분", "역할", "주요 경로"],
        [
            ["사용자", "예측 게임·포인트·출석·게시판", "/prediction, /home, /login"],
            ["쇼핑몰", "스포츠 용품 몰 (계정 공유)", "/shop"],
            ["운영자", "현장 예측 시작·중지·결과", "/manager/*"],
            ["관리자", "경기·회원·몰·모니터링", "/admin/*"],
        ],
    )

    # 2. 전체 구조
    add_heading(doc, "2. 전체 시스템 구조", 1)
    add_para(doc, "클라이언트 → ppamong.com (Express) → MongoDB Atlas / Redis / API-SPORTS", bold=True)
    add_bullets(
        doc,
        [
            "사용자·운영자·관리자·몰 UI는 모두 동일 서버의 SPA/라우트로 제공",
            "실시간: WebSocket (/ws/match) + API-SPORTS 폴링(~2.5초)",
            "인증: JWT (유저·관리자 분리), OAuth(카카오/구글/애플) 지원",
            "도메인: 가비아 DNS → Replit Custom Domain (ppamong.com)",
        ],
    )

    add_heading(doc, "2.1 외부 연동", 2)
    add_table(
        doc,
        ["시스템", "용도", "비고"],
        [
            ["API-SPORTS Baseball", "KBO 일정·이닝·점수·종료", "Pro 플랜 권장 (2026 시즌)"],
            ["MongoDB Atlas", "회원·경기·배팅·몰 데이터", "Replica Set"],
            ["Redis", "세션/실시간 보조", "Replit 내장 가능"],
            ["AdMob", "전면·보상·배너 광고", "공수교대·타자교체"],
        ],
    )

    # 3. 디렉터리
    add_heading(doc, "3. 코드·저장소 구조", 1)
    add_table(
        doc,
        ["경로", "설명"],
        [
            ["client/", "React UI (UserApp / ManagerApp / AdminApp / Mall)"],
            ["server/", "Express API · liveMatch · apiSports · mall"],
            ["shared/", "공통 타입·배당 상수 (predictionOdds 등)"],
            ["docs/", "배포·구조·매뉴얼 문서"],
        ],
    )

    # 4. 예측 게임
    add_heading(doc, "4. 예측 게임", 1)
    add_heading(doc, "4.1 타석 예측 (라운드)", 2)
    add_para(
        doc,
        "운영자가 예측을 시작하면 사용자는 아웃·1루·2루·3루·홈런 중 하나를 선택하고 "
        "포인트를 배팅합니다. API-SPORTS는 타석 결과를 제공하지 않으므로 결과 입력은 운영자 역할입니다.",
    )
    add_table(
        doc,
        ["예측", "배당", "100P 적중 시"],
        [
            ["아웃", "1.2배", "120P"],
            ["1루", "1.5배", "150P"],
            ["2루", "3배", "300P"],
            ["3루", "10배", "1000P"],
            ["홈런", "5배", "500P"],
        ],
    )
    add_bullets(
        doc,
        [
            "배팅 금액: 50 / 100 / 200 / 500 / 1000P",
            "지급식: floor(배팅 × 배당), 미적중 시 배팅분 소멸",
            "같은 라운드에서는 결과 확정 전 선택만 변경 가능(추가 차감 없음)",
        ],
    )

    add_heading(doc, "4.2 승리팀 · 최종 스코어 (사이드 배팅)", 2)
    add_para(
        doc,
        "타석 예측과 별도로, 경기당 1회씩 승리팀·최종 스코어를 맞출 수 있습니다. "
        "1회(이닝) 시작 시 자동 마감되며, 경기 종료 시 API 최종 점수로 자동 정산합니다.",
    )
    add_table(
        doc,
        ["게임", "배당", "배팅 금액", "100P 적중"],
        [
            ["승리팀 (홈/원정)", "2.0배", "100·200·500·1000P만", "200P"],
            ["최종 스코어 정확", "20배", "동일 (100 단위)", "2000P"],
        ],
    )
    add_bullets(
        doc,
        [
            "저장: MatchSideBet 컬렉션 (userId+matchId+type 유니크)",
            "마감: Match.sideBetsLocked (이닝 시작 감지)",
            "무승부·경기 취소: 해당 배팅 환불",
            "사용자 UI: 「홈팀」「원정팀」만 표시 (구단명 미표시)",
            "관리자/운영자: 스코어보드·모니터링에 구단명 표시 가능",
        ],
    )

    add_heading(doc, "4.3 광고·보상", 2)
    add_bullets(
        doc,
        [
            "공수교대(이닝 변경): 전면/보상 광고 제안, 완료 시 500P",
            "너무 일찍 닫으면 보상 없음",
            "타자/라운드 전환 시 하단 배너 광고 가능",
        ],
    )

    # 5. API-SPORTS
    add_heading(doc, "5. API-SPORTS 연동", 1)
    add_para(
        doc,
        "베이스볼 API(v1.baseball.api-sports.io)로 KBO 리그(기본 league id=5) 일정을 조회합니다. "
        "league+date 요청 시 season 파라미터가 필수이며, 날짜 연도(또는 API_SPORTS_SEASON)를 사용합니다.",
    )
    add_table(
        doc,
        ["Secrets", "설명"],
        [
            ["API_SPORTS_KEY", "Dashboard API Key (Pro 권장)"],
            ["API_SPORTS_KBO_LEAGUE_ID", "기본 5"],
            ["API_SPORTS_SEASON", "선택. 미설정 시 요청 날짜의 연도"],
        ],
    )
    add_bullets(
        doc,
        [
            "제공: 일정·이닝·점수·종료 상태",
            "미제공: 타석/플레이바이플레이 텍스트 → 운영자 결과 입력 필요",
            "헬스: GET /api/api-sports/health (healthy, lastError, apiKeyConfigured)",
            "Free 플랜은 최근 시즌(예: 2026) 접근이 막힐 수 있음 → Pro 필요",
        ],
    )

    # 6. 경기 관리
    add_heading(doc, "6. 관리자 경기 관리 (달력)", 1)
    add_para(
        doc,
        "수기 경기 등록 없이, 달력에서 날짜를 클릭하면 해당일 KBO 일정을 API에서 읽어 "
        "DB에 자동 저장하고 API 경기에 연결합니다. (하루 최대 5경기)",
    )
    add_heading(doc, "6.1 사용 흐름", 2)
    add_bullets(
        doc,
        [
            "관리자 → 경기 · 회원 → 경기 관리",
            "달력에서 날짜 클릭 (또는 「오늘 날짜 열기」)",
            "모달 표 표시 + 동시에 sync-from-api-sports 실행",
            "1경기~N경기 생성/갱신, apiSportsGameId·팀명(관리자용)·스코어보드 저장",
            "운영자 op1~op5 배정 자동 동기화",
            "모달에서 「다시 동기화」로 강제 재조회 가능",
            "표의 「모니터링」으로 실시간 게임 모니터링 이동",
        ],
    )
    add_heading(doc, "6.2 모달 표 컬럼", 2)
    add_para(doc, "시간 · 경기 · 원정 · 스코어 · 홈 · 상태 · API 연결 · 모니터링")
    add_heading(doc, "6.3 모니터링", 2)
    add_bullets(
        doc,
        [
            "API 헬스 Green/Red, 스코어보드, 타석 배팅 분포",
            "승리팀·스코어 사이드 배팅 참여 요약",
            "비상 시 수동 제어 전환 (controlMode=manual)",
        ],
    )

    # 7. 역할별
    add_heading(doc, "7. 역할별 운영", 1)
    add_heading(doc, "7.1 관리자", 2)
    add_bullets(
        doc,
        [
            "달력으로 당일/특정일 일정 자동 반영",
            "헬스·배팅·사이드배팅 모니터링",
            "비상 수동 모드, 회원·몰·공지 관리",
        ],
    )
    add_heading(doc, "7.2 운영자", 2)
    add_bullets(
        doc,
        [
            "배정 경기 입장, API 스코어보드 확인",
            "타석 예측 시작 → 유저 배팅 → 중지 → 결과(아웃~홈런)",
            "다음 라운드, 광고 시작/중지",
            "경기 최종 종료는 자동 모드에서 API 종료 상태 시 처리",
        ],
    )
    add_heading(doc, "7.3 사용자", 2)
    add_bullets(
        doc,
        [
            "로고 → 홈(사용설명서·시뮬레이션) / 헤더 「쇼핑몰」→ /shop",
            "경기 참여 → 승리팀·스코어(마감 전) + 타석 예측",
            "하단: 초대·출석·게시·추가(포인트)·로그아웃",
            "사용설명서 /home/guide, 시뮬레이션 /home/simulation",
        ],
    )

    # 8. 주요 API
    add_heading(doc, "8. 주요 API (요약)", 1)
    add_table(
        doc,
        ["메서드", "경로", "설명"],
        [
            ["POST", "/api/admin/matches/sync-from-api-sports", "해당일 일정 DB 저장·연결(최대 5)"],
            ["GET", "/api/api-sports/health", "API 연동 헬스"],
            ["GET", "/api/matches/:id/scoreboard", "라이브 스코어보드"],
            ["POST", "/api/live-match/predictions", "타석 예측 접수"],
            ["POST", "/api/live-match/side-bets", "승리팀/스코어 배팅"],
            ["GET", "/api/live-match/matches/:id/side-bets/me", "내 사이드 배팅"],
            ["GET", "/api/live-match/matches/:id/side-bets/summary", "관리자 요약"],
            ["POST", "/api/live-match/control/:id/end", "경기 종료(+사이드 정산)"],
        ],
    )

    # 9. 데이터 모델
    add_heading(doc, "9. 핵심 데이터", 1)
    add_table(
        doc,
        ["모델", "핵심 필드"],
        [
            ["Match", "name(1~5경기), matchDate, apiSportsGameId, liveScoreboard, sideBetsLocked, controlMode"],
            ["Prediction", "userId, matchId, roundNumber, prediction, amount, status, wonAmount"],
            ["MatchSideBet", "type(winner|score), winnerPick, home/awayScorePick, odds, status"],
            ["User", "points, inviteCode, OAuth 정보"],
            ["Stadium", "구장 (API자동 구장 포함 가능)"],
        ],
    )

    # 10. 배포
    add_heading(doc, "10. 배포·운영", 1)
    add_bullets(
        doc,
        [
            "코드: GitHub main push → Replit Shell에서 git fetch && git reset --hard origin/main",
            "반드시 Deploy(Autoscale) Redeploy (Run만으로는 라이브 미반영)",
            "Secrets: MONGODB_URI, JWT_*, API_SPORTS_KEY, API_SPORTS_KBO_LEAGUE_ID",
            "확인: https://ppamong.com/api/api-sports/health → healthy true",
            "참고: docs/PPAMONG_DEPLOY_CHECKLIST.md, docs/PPAMONG_가비아_DNS_설정.md",
        ],
    )

    # 11. 최근 변경
    add_heading(doc, "11. 최근 주요 변경 (2026-07)", 1)
    add_bullets(
        doc,
        [
            "고정 배당 타석 예측 + API-SPORTS 스코어 동기화",
            "API season 파라미터 필수 대응",
            "사용자·운영자 사용설명서·시뮬레이션",
            "승리팀(2배)·최종스코어(20배) 사이드 배팅, 100P 단위, 1회 시작 마감",
            "사용자 화면 구단명 숨김(홈팀/원정팀)",
            "경기 관리 달력 UI: 날짜 조회 시 API→DB 자동 저장·연결 (최대 5경기)",
        ],
    )

    # 12. 제약
    add_heading(doc, "12. 제약·주의사항", 1)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["타석 결과", "API 없음 → 운영자 수동 결과 필수"],
            ["하루 경기 수", "KBO·시스템 모두 최대 5경기 기준"],
            ["플랜", "Free는 최신 시즌 제한 가능 → Pro 권장"],
            ["구단명", "사용자 UI 미표시(상표·표기 보수 운영)"],
            ["자동 종료", "열린 타석 라운드 미정산 가능 → 종료 전 결과 입력 권장"],
        ],
    )

    add_heading(doc, "13. 관련 문서", 1)
    add_bullets(
        doc,
        [
            "docs/PPAMONG_시스템_구조도.md",
            "docs/PPAMONG_예측게임_시스템흐름.md",
            "docs/PPAMONG_프로젝트_구조.md",
            "docs/PPAMONG_DEPLOY_CHECKLIST.md",
            "docs/ADMIN_MANUAL.md / USER_MANUAL.md / MANAGER_MANUAL.md",
        ],
    )

    footer = doc.add_paragraph()
    run = footer.add_run("— 끝 —")
    set_run_font(run, size=10, color=RGBColor(0x88, 0x88, 0x88))
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
